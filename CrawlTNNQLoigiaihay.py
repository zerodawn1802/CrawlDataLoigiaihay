# Crawl thành ngữ nói quá Loigiaihay
import requests
from bs4 import BeautifulSoup, NavigableString
import os
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches
import time, random

# Hàm thêm văn bản với định dạng vào tài liệu Word
def add_formatted_text(paragraph, text, bold=False, italic=False, underline=False):
    if paragraph.runs and not paragraph.text.startswith(' '):
        paragraph.add_run(' ')
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.underline = underline

# Hàm xử lý nội dung đệ quy (bao gồm văn bản định dạng, bảng và ảnh)
def process_element(element, doc, title, image_dir, headers, paragraph, bold=False, italic=False, underline=False):
    if not paragraph:
        paragraph = doc.add_paragraph()
    if isinstance(element, NavigableString):  # Văn bản thô
        text = element.strip()
        if text:
            add_formatted_text(paragraph, text)
    elif element.name == 'img':  # Ảnh
        img_url = element.get('src')
        if img_url:
            print(f"Tìm thấy ảnh: {img_url}")
            try:
                img_response = requests.get(img_url, headers=headers, stream=True)
                if img_response.status_code == 200:
                    img_filename = os.path.join(image_dir, f"{title[:20]}_{os.path.basename(img_url)}")
                    with open(img_filename, 'wb') as img_file:
                        img_file.write(img_response.content)
                    doc.add_picture(img_filename, width=Inches(5))
                else:
                    print(f"Không thể tải ảnh {img_url}. Mã lỗi: {img_response.status_code}")
            except Exception as e:
                print(f"Lỗi khi tải ảnh {img_url}: {e}")
    elif element.name == 'table':  # Bảng
        rows = element.find_all('tr')
        if rows:
            first_row_cells = rows[0].find_all(['td', 'th'])
            num_cols = len(first_row_cells)
            doc_table = doc.add_table(rows=len(rows), cols=num_cols)
            doc_table.style = 'Table Grid'
            for i, row in enumerate(rows):
                cells = row.find_all(['td', 'th'])
                for j, cell in enumerate(cells):
                    if j < num_cols:
                        doc_table.rows[i].cells[j].text = cell.text.strip()
    elif element.name:
        # Cập nhật trạng thái định dạng dựa trên thẻ hiện tại
        new_bold = bold or element.name in ['strong', 'b']
        new_italic = italic or element.name in ['em', 'i']
        new_underline = underline or element.name == 'u'
        
        for child in element.children:
            if isinstance(child, NavigableString):
                text = child.strip()
                if text:
                    add_formatted_text(paragraph, text, new_bold, new_italic, new_underline)
            else:
                # Truyền trạng thái định dạng mới vào đệ quy
                process_element(child, doc, title, image_dir, headers, paragraph, new_bold, new_italic, new_underline)

def crawl_with_url(url, output_dir="articles"):
    headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'}
    responses = requests.get(url, headers=headers)

    if responses.status_code == 200:
        soups = BeautifulSoup(responses.text, 'html.parser')

        title = soups.find('head').find('title').text.strip()
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)
        
        image_dir = os.path.join(output_dir, "images")
        if not os.path.exists(image_dir):
            os.makedirs(image_dir)
        
        filename = f"{output_dir}/{title[:50].replace('/', '-')}.docx"
        doc = Document()
        list_vb = []
        
        doc.add_heading(f"{title}", level=1)
        box_contents = soups.find('div', id='box-content')
        links = [a.get('href') for a in box_contents.find_all('a')]
        count = 1
        for link in links:
            if not link.startswith("http"):
                link = "https://loigiaihay.com" + link
            response = requests.get(link, headers=headers)
            if response.status_code == 200:
                soup = BeautifulSoup(response.text, 'html.parser')
                box_content = soup.find('div', id='box-content')
                h2_tag = box_content.find('h2', {'dir': 'ltr'})
                if h2_tag:
                    headd = h2_tag.get_text()
                    list_vb.append(headd)
                paragraph = doc.add_paragraph()
                for element in box_content.children:
                    if element.name == 'h2':
                        paragraph = doc.add_paragraph()
                        text = f"{count}. " + element.get_text()
                        add_formatted_text(paragraph, text, bold=True, italic=False, underline=False)
                        count += 1
                        continue
                    if element.name in ['p', 'div']:
                        paragraph = doc.add_paragraph()
                    # Truyền định dạng mặc định
                    process_element(element, doc, title, image_dir, headers, paragraph, bold=False, italic=False, underline=False)
                
                doc.add_paragraph("")
        doc.save(filename)

url = "https://loigiaihay.com/tong-hop-thanh-ngu-su-dung-bien-phap-noi-qua-e36426.html"
crawl_with_url(url)