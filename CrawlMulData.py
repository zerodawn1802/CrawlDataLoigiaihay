import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import urllib.request
import os

# URL của trang web
url = "https://loigiaihay.com/ly-thuyet-vecto-va-cac-phep-toan-vecto-trong-khong-gian-toan-12-canh-dieu-a170112.html"

# Gửi yêu cầu HTTP tới trang web
response = requests.get(url)

# Hàm thêm định dạng cho đoạn văn trong docx
def add_formatted_text(paragraph, text, bold=False, italic=False, underline=False):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.underline = underline

# Hàm tải ảnh và chèn vào docx
def add_image_to_doc(doc, img_url, width=None):
    try:
        # Tải ảnh về tạm thời
        img_filename = "temp_image.jpg"
        urllib.request.urlretrieve(img_url, img_filename)
        # Chèn ảnh vào tài liệu
        paragraph = doc.add_paragraph()
        run = paragraph.add_run()
        run.add_picture(img_filename, width=width)
        # Xóa file tạm sau khi chèn
        os.remove(img_filename)
    except Exception as e:
        print(f"Lỗi khi tải hoặc chèn ảnh {img_url}: {e}")

# Hàm xử lý nội dung HTML và chuyển sang docx
def html_to_docx(doc, element):
    for child in element.children:
        if child.name == 'p':
            paragraph = doc.add_paragraph()
            for subchild in child.descendants:
                if isinstance(subchild, str) and subchild.strip():
                    add_formatted_text(paragraph, subchild.strip())
                elif subchild.name == 'strong' or subchild.name == 'b':
                    add_formatted_text(paragraph, subchild.get_text(strip=True), bold=True)
                elif subchild.name == 'em' or subchild.name == 'i':
                    add_formatted_text(paragraph, subchild.get_text(strip=True), italic=True)
                elif subchild.name == 'u':
                    add_formatted_text(paragraph, subchild.get_text(strip=True), underline=True)
                elif subchild.name == 'img':
                    src = subchild.get('src', '')
                    if src.startswith('//'):
                        src = 'https:' + src  # Thêm giao thức nếu thiếu
                    elif src.startswith('/'):
                        src = 'https://loigiaihay.com' + src  # Thêm domain nếu là đường dẫn tương đối
                    add_image_to_doc(doc, src)
        elif child.name in ['h1', 'h2', 'h3']:
            level = {'h1': 1, 'h2': 2, 'h3': 3}.get(child.name)
            doc.add_heading(child.get_text(strip=True), level=level)
        elif child.name == 'img':
            src = child.get('src', '')
            if src.startswith('//'):
                src = 'https:' + src
            elif src.startswith('/'):
                src = 'https://loigiaihay.com' + src
            add_image_to_doc(doc, src)
        elif child.name == 'div' and child.get('class') != ['Choose-fast']:
            html_to_docx(doc, child)  # Đệ quy xử lý các thẻ div khác
        elif isinstance(child, str) and child.strip():
            doc.add_paragraph(child.strip())

# Kiểm tra nếu yêu cầu thành công
if response.status_code == 200:
    # Parse HTML bằng BeautifulSoup
    soup = BeautifulSoup(response.content, 'html.parser')
    
    # Tìm thẻ div với class "box-content"
    box_content = soup.find('div', id='box-content')
    
    if box_content:
        # Tạo tài liệu docx
        doc = Document()
        doc.add_heading('Nội dung bài học', level=1)
        
        # Chuyển nội dung sang docx
        html_to_docx(doc, box_content)
        
        # Lưu file docx
        output_file = 'noi_dung_bai_hoc.docx'
        doc.save(output_file)
        print(f"Đã lưu nội dung vào file '{output_file}'")
    else:
        print("Không tìm thấy thẻ div với class 'box-content'")
else:
    print(f"Lỗi khi truy cập trang web: {response.status_code}")