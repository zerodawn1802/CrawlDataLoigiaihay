# Crarl dữ liệu không có dạng MathJax
import requests
from bs4 import BeautifulSoup, NavigableString
import os
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import urllib.request
from docx.shared import Inches

# Hàm thêm văn bản với định dạng vào tài liệu Word
def add_formatted_text(paragraph, text, bold=False, italic=False, underline=False):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.underline = underline

# Hàm xử lý nội dung đệ quy (bao gồm văn bản định dạng, bảng và ảnh)
def process_element(element, doc, title, div_id, image_dir, headers, paragraph, bold=False, italic=False, underline=False):
    if isinstance(element, NavigableString):  # Văn bản thô
        text = element.strip()
        if text:
            add_formatted_text(paragraph, text)
    elif element.name == 'img':  # Ảnh
        img_url = element.get('src')
        if img_url:
            print(f"Tìm thấy ảnh: {img_url}")
            try:
                img_response = requests.get(img_url, headers=headers, stream=True)
                if img_response.status_code == 200:
                    img_filename = os.path.join(image_dir, f"{title[:20]}_{div_id}_{os.path.basename(img_url)}")
                    with open(img_filename, 'wb') as img_file:
                        img_file.write(img_response.content)
                    doc.add_picture(img_filename, width=Inches(6))
                else:
                    print(f"Không thể tải ảnh {img_url}. Mã lỗi: {img_response.status_code}")
            except Exception as e:
                print(f"Lỗi khi tải ảnh {img_url}: {e}")
    elif element.name == 'table':  # Bảng
        rows = element.find_all('tr')
        if rows:
            first_row_cells = rows[0].find_all(['td', 'th'])
            num_cols = len(first_row_cells)
            doc_table = doc.add_table(rows=len(rows), cols=num_cols)
            doc_table.style = 'Table Grid'
            for i, row in enumerate(rows):
                cells = row.find_all(['td', 'th'])
                for j, cell in enumerate(cells):
                    if j < num_cols:
                        doc_table.rows[i].cells[j].text = cell.text.strip()
    elif element.name:
        # Cập nhật trạng thái định dạng dựa trên thẻ hiện tại
        new_bold = bold or element.name in ['strong', 'b']
        new_italic = italic or element.name in ['em', 'i']
        new_underline = underline or element.name == 'u'
        
        for child in element.children:
            if isinstance(child, NavigableString):
                text = child.strip()
                if text:
                    add_formatted_text(paragraph, text, new_bold, new_italic, new_underline)
            else:
                # Truyền trạng thái định dạng mới vào đệ quy
                process_element(child, doc, title, image_dir, headers, paragraph, new_bold, new_italic, new_underline)

# Hàm crawl nội dung một trang và lưu dưới dạng .docx
def crawl_single_page(url, output_dir="articles"):
    headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'}
    response = requests.get(url, headers=headers)
    
    if response.status_code == 200:
        soup = BeautifulSoup(response.text, 'html.parser')
        
        # Lấy tiêu đề
        title = soup.find('h1').text.strip()
        print(f"Đang crawl: {title}")
        
        # Tìm tất cả các div với class="box-question top20"
        content_divs = soup.find_all('div', class_='box-question top20')
        if not content_divs:
            print(f"Không tìm thấy div với class='box-question top20' tại {url}")
            return
        
        # Tạo thư mục nếu chưa có
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)
        
        # Thư mục tạm để lưu ảnh
        image_dir = os.path.join(output_dir, "images")
        if not os.path.exists(image_dir):
            os.makedirs(image_dir)
        
        # Tạo file .docx
        filename = f"{output_dir}/{title[:50].replace('/', '-')}.docx"
        doc = Document()
        
        # Thêm tiêu đề vào tài liệu
        doc.add_heading(f"Tiêu đề: {title}", level=1)
        
        # Lặp qua từng div
        for div in content_divs:
            div_id = div.get('id', 'Không có ID')

            # Loại bỏ tất cả các div với class="section-explanation-tab"
            for explanation in div.find_all('div', class_='section-explanation-tab'):
                explanation.decompose()
            for explanation in div.find_all('div', class_='Choose-fast'):
                explanation.decompose()

            
            # Xử lý toàn bộ nội dung trong div
            paragraph = doc.add_paragraph()  # Tạo một paragraph ban đầu cho mỗi div
            for element in div.children:
                if element.name in ['p', 'div']:  # Tạo paragraph mới cho <p> hoặc <div>
                    paragraph = doc.add_paragraph()
                process_element(element, doc, title, div_id, image_dir, headers, paragraph, bold=False, italic=False, underline=False)
            
            # Thêm dòng trống sau mỗi phần (chỉ một lần)
            doc.add_paragraph("")
        
        # Lưu file .docx
        doc.save(filename)
        print(f"Đã lưu: {filename}")
    else:
        print(f"Không thể truy cập {url}. Mã lỗi: {response.status_code}")

# URL của trang muốn crawl
url = "https://loigiaihay.com/-giai-muc-1-trang-567-sgk-toan-12-tap-1-ket-noi-tri-thuc-a157122.html"  # Thay bằng URL của bạn
crawl_single_page(url)