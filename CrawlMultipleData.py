import requests
from bs4 import BeautifulSoup
import time
import os
from docx import Document

# Hàm crawl nội dung một bài
def crawl_single_page(url, output_dir="articles"):
    headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'}
    response = requests.get(url, headers=headers)
    
    if response.status_code == 200:
        soup = BeautifulSoup(response.text, 'html.parser')
        
        # Lấy tiêu đề
        title = soup.find('h1').text.strip()
        print(f"Đang crawl: {title}")
        
        # Tìm tất cả các div với class="box-question top20"
        content_divs = soup.find_all('div', class_='box-question top20')
        if not content_divs:
            print(f"Không tìm thấy div với class='box-question top20' tại {url}")
            return
        
        # Tạo thư mục nếu chưa có
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)
        
        # Tạo file .docx
        filename = f"{output_dir}/{title[:50].replace('/', '-')}.docx"
        doc = Document()
        
        # Thêm tiêu đề vào tài liệu
        doc.add_heading(f"Tiêu đề: {title}", level=1)
        
        # Lặp qua từng div
        for div in content_divs:
            
            # Kiểm tra xem div có chứa bảng không
            table = div.find('table')
            if table:
                # Lấy tất cả các hàng trong bảng
                rows = table.find_all('tr')
                if rows:
                    # Đếm số cột dựa trên hàng đầu tiên
                    first_row_cells = rows[0].find_all(['td', 'th'])
                    num_cols = len(first_row_cells)
                    
                    # Tạo bảng trong tài liệu Word
                    doc_table = doc.add_table(rows=len(rows), cols=num_cols)
                    doc_table.style = 'Table Grid'  # Thêm đường viền cho bảng
                    
                    # Điền dữ liệu vào bảng
                    for i, row in enumerate(rows):
                        cells = row.find_all(['td', 'th'])
                        for j, cell in enumerate(cells):
                            if j < num_cols:  # Đảm bảo không vượt quá số cột
                                doc_table.rows[i].cells[j].text = cell.text.strip()
                doc.add_paragraph("")  # Thêm dòng trống sau bảng
            else:
                # Nếu không có bảng, lấy các đoạn văn hoặc văn bản thô
                paragraphs = div.find_all('p')
                if paragraphs:
                    for i, p in enumerate(paragraphs):
                        doc.add_paragraph(f"{p.text.strip()}")
                    doc.add_paragraph("")
                else:
                    doc.add_paragraph(f"{div.text.strip()}")
                    doc.add_paragraph("")
        
        # Lưu file .docx
        doc.save(filename)
        print(f"Đã lưu: {filename}")
    else:
        print(f"Không thể truy cập {url}. Mã lỗi: {response.status_code}")

# Hàm crawl nhiều bài từ trang danh mục (giữ nguyên)
def crawl_multiple_articles(start_url, max_articles=10):
    headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'}
    visited = set()
    to_crawl = [start_url]
    articles_crawled = 0
    
    while to_crawl and articles_crawled < max_articles:
        url = to_crawl.pop(0)
        if url in visited:
            continue
        
        try:
            response = requests.get(url, headers=headers)
            if response.status_code == 200:
                soup = BeautifulSoup(response.text, 'html.parser')
                
                # Tìm tất cả liên kết bài viết
                for link in soup.select('a[href*="/soan-bai-"]'):  # Giữ nguyên hoặc điều chỉnh
                    article_url = link['href']
                    if not article_url.startswith('http'):
                        article_url = "https://loigiaihay.com" + article_url
                    
                    if article_url not in visited:
                        crawl_single_page(article_url)
                        visited.add(article_url)
                        articles_crawled += 1
                        if articles_crawled >= max_articles:
                            break
                        time.sleep(1)
                
                # Tìm trang tiếp theo (nếu có phân trang)
                next_page = soup.find('a', class_='next')
                if next_page and next_page['href'] not in visited:
                    next_url = next_page['href']
                    if not next_url.startswith('http'):
                        next_url = "https://loigiaihay.com" + next_url
                    to_crawl.append(next_url)
                
                visited.add(url)
            time.sleep(1)
            
        except requests.RequestException as e:
            print(f"Lỗi khi crawl {url}: {e}")
    
    print(f"\nĐã crawl {articles_crawled} bài")

# Chạy chương trình
start_url = "https://loigiaihay.com/soan-van-12-ket-noi-tri-thuc-c56.html"
crawl_multiple_articles(start_url, max_articles=5)  # Crawl 5 bài để thử nghiệm