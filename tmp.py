# Crawl thành ngữ nói quá Loigiaihay
import requests
from bs4 import BeautifulSoup, NavigableString
import os
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches

# Hàm thêm văn bản với định dạng vào tài liệu Word
def add_formatted_text(paragraph, text, bold=False, italic=False, underline=False):
    # if paragraph.runs and not paragraph.text.endswith(' '):
    #     paragraph.add_run(' ')
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.underline = underline

def int_to_roman(n):
    roman_dict = {
        1: "I", 4: "IV", 5: "V", 9: "IX", 10: "X",
        40: "XL", 50: "L", 90: "XC", 100: "C",
        400: "CD", 500: "D", 900: "CM", 1000: "M"
    }
    roman_value = ""
    for value in sorted(roman_dict.keys(), reverse=True):
        while n >= value:
            roman_value += roman_dict[value]
            n -= value
    return roman_value

# Hàm xử lý nội dung đệ quy (bao gồm văn bản định dạng, bảng và ảnh)
def process_element(element, doc, title, image_dir, headers, paragraph, bold=False, italic=False, underline=False):
    # if not paragraph:
    #     paragraph = doc.add_paragraph()
    if isinstance(element, NavigableString):  # Văn bản thô
        text = element.strip()
        if text:
            add_formatted_text(paragraph, text)
    elif element.name == 'img':  # Ảnh
        img_url = element.get('src')
        if img_url:
            print(f"Tìm thấy ảnh: {img_url}")
            try:
                img_response = requests.get(img_url, headers=headers, stream=True)
                if img_response.status_code == 200:
                    img_filename = os.path.join(image_dir, f"{title[:20]}_{os.path.basename(img_url)}")
                    with open(img_filename, 'wb') as img_file:
                        img_file.write(img_response.content)
                    doc.add_picture(img_filename, width=Inches(5))
                else:
                    print(f"Không thể tải ảnh {img_url}. Mã lỗi: {img_response.status_code}")
            except Exception as e:
                print(f"Lỗi khi tải ảnh {img_url}: {e}")
    elif element.name == 'table':  # Bảng
        rows = element.find_all('tr')
        if rows:
            first_row_cells = rows[0].find_all(['td', 'th'])
            num_cols = len(first_row_cells)
            doc_table = doc.add_table(rows=len(rows), cols=num_cols)
            doc_table.style = 'Table Grid'
            for i, row in enumerate(rows):
                cells = row.find_all(['td', 'th'])
                for j, cell in enumerate(cells):
                    if j < num_cols:
                        doc_table.rows[i].cells[j].text = cell.text.strip()
    elif element.name == 'ul':  # Xử lý danh sách <ul>
        for li in element.find_all('li', recursive=False):
            paragraph = doc.add_paragraph()  # Tạo paragraph mới cho mỗi <li>
            process_element(li, doc, title, image_dir, headers, paragraph, bold, italic, underline)
    
    elif element.name in ['p', 'div', 'span']:  # Xử lý <p>, <div>
        # Chỉ tạo paragraph mới nếu đây là thẻ cấp cao và không phải con của <ul>
        if element.parent.name not in ['ul', 'li']:
            paragraph = doc.add_paragraph()
        for child in element.children:
            if isinstance(child, NavigableString):
                text = child.strip()
                if text:
                    add_formatted_text(paragraph, text, bold, italic, underline)
            else:
                new_bold = bold or child.name in ['strong', 'b']
                new_italic = italic or child.name in ['em', 'i']
                new_underline = underline or child.name == 'u'
                process_element(child, doc, title, image_dir, headers, paragraph, new_bold, new_italic, new_underline)
    elif element.name:  # Các thẻ khác (p, div, span, strong, em, u, v.v.)
        bold = element.name in ['strong', 'b']
        italic = element.name in ['em', 'i']
        underline = element.name == 'u'
        
        for child in element.children:
            if isinstance(child, NavigableString):
                text = child.strip()
                if text:
                    add_formatted_text(paragraph, text, bold, italic, underline)
                    if paragraph.runs and not paragraph.text.endswith(' '):
                        paragraph.add_run(' ')
            else:
                process_element(child, doc, title, image_dir, headers, paragraph)

def crawl_with_url(url, output_dir="articles"):
    headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'}
    responses = requests.get(url, headers=headers)

    if responses.status_code == 200:
        soups = BeautifulSoup(responses.text, 'html.parser')

        title = soups.find('head').find('title').text.strip()
        # Tạo thư mục nếu chưa có
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)
        
        # Thư mục tạm để lưu ảnh
        image_dir = os.path.join(output_dir, "images")
        if not os.path.exists(image_dir):
            os.makedirs(image_dir)
        
        # Tạo file .docx
        filename = f"{output_dir}/{title[:100].replace('/', '-')}.docx"
        list_vb = []
        doc = Document()
        title_paragraph = soups.find('div', attrs={
            'style': "color:rgb(57 117 234);font-family: 'OpenSansBold';margin: 5px 0 16px 0; font-size: 24px;font-weight: bold;text-align: center"
        }).text.strip()
        paragraph = doc.add_paragraph()
        doc.add_paragraph(add_formatted_text(paragraph, title_paragraph, bold=True, italic=False, underline=False))
        box_contents = soups.find('div', id='box-content')
        count = 1
        for element in box_contents.children:
            if element.name == 'script' or element.__class__== 'wiki-header-search':
                continue
            h2_tag = element.find('h2')
            if h2_tag != None and h2_tag != -1:
                n = int_to_roman(int(h2_tag.get_text().strip().split('.')[0])) + "." + h2_tag.get_text().strip().split('.')[1]
                paragraph = doc.add_paragraph()
                add_formatted_text(paragraph, n, bold=True, italic=False, underline=False)
                count = 1
            if element.name == 'a':
                tmp_link = element.get('href')
                if not tmp_link.startswith("http"):
                    tmp_link = "https://loigiaihay.com" + tmp_link
                    header = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'}
                    response = requests.get(tmp_link, headers=header)
                    if response.status_code == 200:
                        soup = BeautifulSoup(response.text, 'html.parser')
                        box_content = soup.find('div', id='box-content')
                        tmp_text = ""
                        first_sp = True
                        check = {}
                        for element in box_content.children:
                            if element.name == 'h2':
                                has_sp = element.find('span')
                                if has_sp:
                                    for p_tag in element.find_all('span', recursive=False):
                                        has_img = p_tag.find('img')
                                        has_p = p_tag.find('span')
                                        if has_p:
                                            for p in p_tag.find_all('span'):
                                                if first_sp == True:
                                                    tmp_text += f"{count}. {p.get_text()}"
                                                    count += 1
                                                    first_sp = False
                                                else:
                                                    tmp_text += f"{p.get_text()}"
                                        elif has_img:
                                            for img_tag in p_tag.find_all('img'):
                                                paragraph = doc.add_paragraph()
                                                process_element(img_tag, doc, title, image_dir, header, paragraph)
                                        elif not has_img and not has_p:
                                            if first_sp == True:
                                                tmp_text += f"{count}. {p_tag.get_text()}"
                                                count += 1
                                                first_sp = False
                                            else:
                                                tmp_text += f"{p_tag.get_text()}"
                                elif not has_sp:
                                    if first_sp == True:
                                        tmp_text += f"{count}. {element.get_text()}"
                                        count += 1
                                        first_sp = False
                                    elif first_sp == False:
                                        tmp_text += f"{element.get_text()}"
                            if tmp_text:
                                paragraph = doc.add_paragraph()
                                add_formatted_text(paragraph, tmp_text, bold=True, italic=False, underline=False)
                                tmp_text = ""
                            if element.name in ['p', 'div', 'ul']:
                                paragraph = doc.add_paragraph()
                                if tmp_text:
                                    add_formatted_text(paragraph, tmp_text, bold=True, italic=False, underline=False)
                                    tmp_text = ""
                                    paragraph = doc.add_paragraph()
                                # Truyền định dạng mặc định
                                process_element(element, doc, title, image_dir, header, paragraph)
                            if element.__class__ == 'accordion-wiki':
                                break
                        doc.add_paragraph("")
            else:
                paragraph = doc.add_paragraph()
                # Truyền định dạng mặc định
                process_element(element, doc, title, image_dir, headers, paragraph)
        doc.save(filename)
                

url = "https://loigiaihay.com/tong-hop-ca-dao-tuc-ngu-ve-thay-co-nhan-ngay-20-11-e36238.html"
crawl_with_url(url)